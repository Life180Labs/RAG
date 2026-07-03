"""Unit tests for each format parser (docs/05-task.md Phase 5), run
against real generated documents (no mocks) — a PDF built with PyMuPDF, a
DOCX built with python-docx, etc."""

import io

import fitz
from docx import Document as DocxDocument

from document_worker.parsing import factory


def test_txt_parser_detects_title_and_paragraphs():
    parse_fn, name = factory.get_parser("txt")
    content = b"My Document Title\n\nFirst paragraph body text.\n\nSecond paragraph body text."

    blocks, page_count = parse_fn(content)

    assert name == "native"
    assert page_count is None
    assert blocks[0] == {"type": "title", "text": "My Document Title", "level": None, "page": None}
    assert blocks[1]["type"] == "paragraph"
    assert blocks[2]["type"] == "paragraph"


def test_markdown_parser_detects_all_block_types():
    parse_fn, _ = factory.get_parser("md")
    content = (
        b"# Title\n\nBody paragraph.\n\n- item one\n- item two\n\n"
        b"```python\nprint(1)\n```\n\n![a photo](img.png)\n"
    )

    blocks, _ = parse_fn(content)
    types = [b["type"] for b in blocks]

    assert types == ["title", "paragraph", "list", "code", "image"]
    assert blocks[2]["text"] == "item one\nitem two"
    assert blocks[4]["text"] == "a photo"


def test_html_parser_detects_all_block_types():
    parse_fn, _ = factory.get_parser("html")
    content = (
        b"<html><body>"
        b"<h1>Page Title</h1>"
        b"<p>First paragraph.</p>"
        b"<ul><li>item a</li><li>item b</li></ul>"
        b"<pre>code line 1</pre>"
        b'<table><tr><th>Name</th><th>Age</th></tr><tr><td>Alice</td><td>30</td></tr></table>'
        b'<img src="x.png" alt="a photo">'
        b"</body></html>"
    )

    blocks, _ = parse_fn(content)
    types = [b["type"] for b in blocks]

    assert types == ["title", "paragraph", "list", "list", "code", "table", "image"]
    assert blocks[5]["text"] == "Name | Age\nAlice | 30"


def test_csv_parser_produces_single_table_block():
    parse_fn, name = factory.get_parser("csv")
    content = b"name,age\nAlice,30\nBob,25\n"

    blocks, _ = parse_fn(content)

    assert name == "pandas"
    assert len(blocks) == 1
    assert blocks[0]["type"] == "table"
    assert blocks[0]["text"] == "name | age\nAlice | 30\nBob | 25"


def test_json_parser_produces_pretty_printed_code_block():
    parse_fn, name = factory.get_parser("json")
    content = b'{"a": 1, "b": [1, 2, 3]}'

    blocks, _ = parse_fn(content)

    assert name == "native"
    assert blocks[0]["type"] == "code"
    assert '"a": 1' in blocks[0]["text"]


def test_xml_parser_extracts_leaf_element_text_as_paragraphs():
    parse_fn, _ = factory.get_parser("xml")
    content = b"<root><title>Doc Title</title><body><para>Hello world</para></body></root>"

    blocks, _ = parse_fn(content)

    assert [b["type"] for b in blocks] == ["paragraph", "paragraph"]
    assert blocks[0]["text"] == "Doc Title"
    assert blocks[1]["text"] == "Hello world"


def test_docx_parser_detects_headings_lists_and_tables():
    document = DocxDocument()
    document.add_heading("My Report", level=0)
    document.add_heading("Section One", level=1)
    document.add_paragraph("A normal body paragraph.")
    document.add_paragraph("Bullet item one", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    buffer = io.BytesIO()
    document.save(buffer)

    parse_fn, name = factory.get_parser("docx")
    blocks, page_count = parse_fn(buffer.getvalue())

    assert name == "python-docx"
    assert page_count is None
    assert blocks[0] == {"type": "title", "text": "My Report", "level": None, "page": None}
    assert blocks[1] == {"type": "heading", "text": "Section One", "level": 1, "page": None}
    assert blocks[2]["type"] == "paragraph"
    assert blocks[3]["type"] == "list"
    assert blocks[4]["type"] == "table"
    assert blocks[4]["text"] == "A | B\n1 | 2"


def test_pdf_parser_detects_heading_paragraph_list_and_code_by_font():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "My PDF Report", fontsize=24)
    page.insert_text((72, 120), "A normal body paragraph in the document.", fontsize=11)
    page.insert_text((72, 150), "- bullet item one", fontsize=11)
    page.insert_text((72, 170), "x = 1  # monospace code line", fontsize=10, fontname="Courier")
    pdf_bytes = doc.tobytes()
    doc.close()

    parse_fn, name = factory.get_parser("pdf")
    blocks, page_count = parse_fn(pdf_bytes)

    assert name == "pymupdf"
    assert page_count == 1
    assert blocks[0]["type"] == "title"
    assert blocks[0]["text"] == "My PDF Report"
    assert blocks[1]["type"] == "paragraph"
    assert blocks[2]["type"] == "list"
    assert blocks[3]["type"] == "code"
    assert all(b["page"] == 1 for b in blocks)
