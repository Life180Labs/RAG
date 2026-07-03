"""DOCX parser (python-docx). Page count is unknown without rendering
(DOCX has no fixed pagination), so `page` is always None here — the
metadata step reports `page_count: None` for this format."""

import io

from docx import Document as DocxDocument


def parse(content: bytes) -> tuple[list[dict], int | None]:
    doc = DocxDocument(io.BytesIO(content))
    blocks: list[dict] = []
    title_assigned = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style = (paragraph.style.name or "").lower() if paragraph.style else ""

        if style == "title":
            blocks.append({"type": "title", "text": text, "level": None, "page": None})
            title_assigned = True
        elif style.startswith("heading"):
            digits = "".join(ch for ch in style if ch.isdigit())
            level = int(digits) if digits else 1
            block_type = "heading" if title_assigned else "title"
            title_assigned = True
            level_value = level if block_type == "heading" else None
            blocks.append(
                {"type": block_type, "text": text, "level": level_value, "page": None}
            )
        elif style.startswith("list") or text.startswith(("-", "*", "•")):
            blocks.append({"type": "list", "text": text, "level": None, "page": None})
        elif "code" in style:
            blocks.append({"type": "code", "text": text, "level": None, "page": None})
        else:
            blocks.append({"type": "paragraph", "text": text, "level": None, "page": None})

    for table in doc.tables:
        rows_text = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        text = "\n".join(rows_text)
        if text.strip():
            blocks.append({"type": "table", "text": text, "level": None, "page": None})

    for _ in doc.inline_shapes:
        blocks.append({"type": "image", "text": "[Image]", "level": None, "page": None})

    return blocks, None
